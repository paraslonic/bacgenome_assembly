# coding=utf-8
import argparse
import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
import re
import sys
from bakgen.docer.parse import parse_read_stats, parse_assembly_stats, parse_bacterial

sys.dont_write_bytecode = True

PICTURE_WIDTH_CM = 15
PICTURE_1 = 'circularGenomeComparison.jpg'
PICTURE_2 = 'CoreGenesMLTree.png'
PICTURE_3 = 'CoreGenesDistanceHeatmap.png'
PICTURE_4 = 'PanGenomeDistanceHeatmap.png'

parser = argparse.ArgumentParser()
parser.add_argument('-d', '--directory', action='store',
                    help='A directory containing all needed files')
args = parser.parse_args()

TABLE_1 = 'readsStats.txt'
TABLE_2 = 'assemblyStats.txt'
bacteria_name = args.directory.split('/')[-1].replace('_', ' ')
bacteria_name_underline = args.directory.split('/')[-1]
TABLE_3 = bacteria_name_underline + '.txt'

print 'docer started'
print 'input directory: ' + args.directory

os.chdir(args.directory)

with open(TABLE_1) as f: table1_lines = f.readlines()
with open(TABLE_2) as f: table2_lines = f.readlines()
with open(TABLE_3) as f: table3_lines = f.readlines()
metabolism_regex = re.compile('Base pathways</td><td align=LEFT valign=TOP>(\d*)')
with open("metabolism.html", "r") as metabolism_file:
    metabolism = re.findall(metabolism_regex, metabolism_file.read().replace('\n', ''))[0]

read_stats = parse_read_stats(table1_lines)
assembly_stats = parse_assembly_stats(table2_lines)
bacterial = parse_bacterial(table3_lines)

document = Document()

style = document.styles['Heading 1']
font = style.font
font.size = Pt(16)
font.bold = True
font.color.rgb = RGBColor(0x0, 0x0, 0x0)
font.name = 'Arial'
style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

style = document.styles['Heading 2']
font = style.font
font.size = Pt(16)
font.bold = True
font.color.rgb = RGBColor(0x0, 0x0, 0x0)
font.name = 'Arial'

style = document.styles['Heading 3']
font = style.font
font.size = Pt(14)
font.bold = True
font.color.rgb = RGBColor(0, 0, 0)
font.name = 'Arial'

style = document.styles['Normal']
font = style.font
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
font.name = 'Arial'

style = document.styles['Normal']
font = style.font
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
font.name = 'Arial'

style = document.styles.add_style('tochange', WD_STYLE_TYPE.PARAGRAPH)
style.font.color.rgb = RGBColor(255, 0, 0)

document.add_heading(u'Отчёт по геномному анализу', level=1)

document.add_heading(u'Содержание', level=2)
document.add_paragraph(u'Содержание\n'
                       u'Введение\n'
                       u'Сборка генома de novo и аннотация\n'
                       u'\tТаксономическая принадлежность образца\n'
                       u'\tИсходные данные\n'
                       u'\tСборка de novo\n'
                       u'\tАннотация\n'
                       u'\tАнализ метаболических путей\n'
                       u'Сравнение с ближайшими организмами\n'
                       u'\tВыравнивание нуклеотидных последовательностей геномов ближайших референсных организмов (представители вида L. plantarum)\n'
                       u'\tАнализ групп гомологии\n'
                       u'\tФилогенетический анализ\n'
                       u'\tВизуализация расстояний между геномами\n'
                       u'\tПоиск факторов вирулентности\n'
                       u'Выводы, примечания\n'
                       u'Методы\n'
                       u'Литература\n')
document.add_page_break()

document.add_heading(u'Введение', level=2)
document.add_paragraph(
    u'Документ содержит отчет по геномному анализу бактерии %s. Была произведена de novo сборка генома (%s контигов). В процессе аннотации и анализа генома был определен %s ген, входящий в состав %s метаболических путей. Для сравнительного геномного анализа исследуемой бактерии и наиболее близких организмов использовались последовательности геномов вида <ВИД БАКТЕРИИ>, публично доступные в базе данных NCBI. Наиболее близким к исследуемому оказался штамм <ШТАММ>, а следующий по ML - <ШТАММ>.  Было построено филогенетическое дерево для всех рассматриваемых штаммов, построены диаграммы сходства геномов по различным метрикам' % (bacteria_name, assembly_stats.assemblies[3].contigs, bacterial.cds, metabolism))

document.add_heading(u'Сборка генома de novo и аннотация', level=2)
document.add_heading(u'Таксономическая принадлежность образца', level=3)
document.add_paragraph(bacteria_name)

document.add_heading(u'Исходные данные', level=3)

table = document.add_table(rows=3, cols=3, style='Table Grid')
table.rows[0].cells[1].text = '\n' + u'Количество ридов'
table.rows[0].cells[2].text = '\n' + u'Средняя длина рида, п.н.'
table.rows[1].cells[0].text = '\n' + read_stats.name1
table.rows[1].cells[1].text = '\n' + read_stats.count1
table.rows[1].cells[2].text = '\n' + read_stats.length1
table.rows[2].cells[0].text = '\n' + read_stats.name2
table.rows[2].cells[1].text = '\n' + read_stats.count2
table.rows[2].cells[2].text = '\n' + read_stats.length2
document.add_paragraph(u'Ожидаемое среднее покрытие составляет %s' % read_stats.coverage)

document.add_page_break()
document.add_heading(u'Сборка de novo', level=3)

table2 = document.add_table(rows=5, cols=6, style='Table Grid')
table2.rows[0].cells[1].text = u'\nКол-во контигов'
table2.rows[0].cells[2].text = u'\nN50'
table2.rows[0].cells[3].text = u'\nДлина сборки'
table2.rows[0].cells[4].text = u'\nGC'
table2.rows[0].cells[5].text = u'\nМакс. длина контига'

table2.rows[1].cells[0].text = '\n' + assembly_stats.assemblies[0].name
table2.rows[1].cells[1].text = '\n' + assembly_stats.assemblies[0].contigs
table2.rows[1].cells[2].text = '\n' + assembly_stats.assemblies[0].n50
table2.rows[1].cells[3].text = '\n' + assembly_stats.assemblies[0].length
table2.rows[1].cells[4].text = '\n' + assembly_stats.assemblies[0].gc
table2.rows[1].cells[5].text = '\n' + assembly_stats.assemblies[0].maximum

table2.rows[2].cells[0].text = '\n' + assembly_stats.assemblies[1].name
table2.rows[2].cells[1].text = '\n' + assembly_stats.assemblies[1].contigs
table2.rows[2].cells[2].text = '\n' + assembly_stats.assemblies[1].n50
table2.rows[2].cells[3].text = '\n' + assembly_stats.assemblies[1].length
table2.rows[2].cells[4].text = '\n' + assembly_stats.assemblies[1].gc
table2.rows[2].cells[5].text = '\n' + assembly_stats.assemblies[1].maximum

table2.rows[3].cells[0].text = '\n' + assembly_stats.assemblies[2].name
table2.rows[3].cells[1].text = '\n' + assembly_stats.assemblies[2].contigs
table2.rows[3].cells[2].text = '\n' + assembly_stats.assemblies[2].n50
table2.rows[3].cells[3].text = '\n' + assembly_stats.assemblies[2].length
table2.rows[3].cells[4].text = '\n' + assembly_stats.assemblies[2].gc
table2.rows[3].cells[5].text = '\n' + assembly_stats.assemblies[2].maximum

table2.rows[4].cells[0].text = '\n' + assembly_stats.assemblies[3].name
table2.rows[4].cells[1].text = '\n' + assembly_stats.assemblies[3].contigs
table2.rows[4].cells[2].text = '\n' + assembly_stats.assemblies[3].n50
table2.rows[4].cells[3].text = '\n' + assembly_stats.assemblies[3].length
table2.rows[4].cells[4].text = '\n' + assembly_stats.assemblies[3].gc
table2.rows[4].cells[5].text = '\n' + assembly_stats.assemblies[3].maximum

document.add_heading(u'Аннотация', level=3)

table3 = document.add_table(rows=5, cols=2, style='Table Grid')
table3.rows[0].cells[0].text = '\n' + u'Тип'
table3.rows[0].cells[1].text = '\n' + u'Количество'
table3.rows[1].cells[0].text = '\n' + 'CDS'
table3.rows[1].cells[1].text = '\n' + bacterial.cds
table3.rows[2].cells[0].text = '\n' + 'tRNA'
table3.rows[2].cells[1].text = '\n' + bacterial.trna
table3.rows[3].cells[0].text = '\n' + 'rRNA'
table3.rows[3].cells[1].text = '\n' + bacterial.rrna
table3.rows[4].cells[0].text = '\n' + 'tmRNA'
table3.rows[4].cells[1].text = '\n' + bacterial.tmrna

document.add_paragraph(
    u'Файл с аннотацией генома в формате GenBank: %s.gbk. \n Файл с транслированными аминокислотными последовательностями генов: %s.faa' %(bacteria_name_underline, bacteria_name_underline))
document.add_heading(u'Анализ метаболических путей', level=3)
document.add_paragraph(u'Количество идентифицированных метаболических путей: %s' % metabolism)
document.add_paragraph(u'Файл с таблицей метаболических путей: metabolism.html')
document.add_paragraph(
    u'(Наиболее информативные поля: Pathway - имя метаболического пути, Total Rxns - количество генов в этом пути, Rxns Present in … - количество генов в этом пути, которые есть в анализируемом геноме)')

document.add_page_break()

document.add_heading(u'Сравнение с ближайшими организмами', level=2)
document.add_heading(
    u'Выравнивание нуклеотидных последовательностей геномов ближайших референсных организмов'
    u' (представители вида L. plantarum)',
    level=3)
document.add_picture(PICTURE_1, width=Cm(PICTURE_WIDTH_CM))

document.add_heading(u'Анализ групп гомологии', level=3)
document.add_paragraph(
    u'Таблица содержит названия групп гомологии и данные об их представленности в рассматриваемых штаммах.')
document.add_paragraph(u'Таблица: ortho_table.xls')

document.add_heading(u'Филогенетический анализ', level=3)
document.add_paragraph(
    u'Дерево, построенное по методу максимального правдоподобия на основе последовательности всех общих (имеющихся у всех рассматриваемых организмов и однокопийных) генов:')
# сначала ругнулся на CoreGenesMLTreeNamed. Какое имя будет в итоге?
document.add_picture(PICTURE_2, width=Cm(PICTURE_WIDTH_CM))
document.add_paragraph(u'Дерево в формате newick (для программы FigTree): CoreGenesMLTreeNamed.nwk')
document.add_paragraph(u'Дерево в формате PDF: CoreGenesMLTreeNamed.pdf')
document.add_paragraph(
    u'Наиболее близким к исследуемому оказался штамм ...', style=document.styles['tochange'])

document.add_heading(u'Визуализация расстояний между геномами', level=3)
document.add_paragraph(
    u'A) по нуклеотидным заменам в последовательностях общих (коровых) генов. Чем выше нуклеотидное сходство последовательностей генов, тем меньше расстояние (краснее цвет). Геномы упорядочены в соответствии с филогенетическим деревом.')
document.add_picture(PICTURE_3, width=Cm(PICTURE_WIDTH_CM))
document.add_paragraph(u'Файл в формате PDF: CoreGenesDistanceHeatmap.pdf')

document.add_paragraph(
    u'B) по генному составу. Чем выше степень совпадения наличия/отсутствия генов, тем меньше расстояние (краснее цвет).')
document.add_picture(PICTURE_4, width=Cm(PICTURE_WIDTH_CM))
document.add_paragraph(u'Файл в формате PDF: PanGenomeDistanceHeatmap.pdf')

document.add_heading(u'Поиск факторов вирулентности', level=3)
document.add_paragraph(
    u'Для поиска факторов вирулентности использовалась база данных VFDB (www.mgc.ac.cn/VFs).')
document.add_paragraph(u'В столбце Name приведены названия генов из VFDB.')
document.add_paragraph(
    u'Обратите внимание на столбцы percIdentity (доля совпадающих аминокислот в выровнявшихся последовательностях), alnLength (длина выравнивания). Чем больше эти значения, тем больше вероятность того, что аннотируемый ген выполняет функцию, указанную в колонке Name.')
document.add_paragraph(u'Таблица: virulence.xls')

document.add_heading(u'Выводы, примечания', level=2)
document.add_paragraph(u'Выводы', style=document.styles['tochange'])
document.add_page_break()

document.add_heading(u'Методы', level=2)
document.add_paragraph(u'1) Сборка генома de novo')
document.add_paragraph(u'Исходные риды были обработаны утилитой Trimmomatic со стандартными параметрами для Illumina. Затем обработанные риды использовались для сборки генома de novo при помощи программ Spades, MIRA 4.0, Newbler 2.6. Наилучшая сборка, использованная для последующего анализа, была выбрана на основании максимального значения N50 и суммарной длины сборки.')
document.add_paragraph(u'Для идентификации и фильтрации загрязнений использовались утилиты BLAST и MEGAN.')

document.add_paragraph(u'2) Аннотация')
document.add_paragraph(u'Аннотация генома производилась с помощью утилиты Prokka v. 1.11 [Seemann, 2014].')

document.add_paragraph(u'3) Анализ метаболических путей')
document.add_paragraph(u'Таблица метаболических путей, реконструированных на основе аннотации генома, была получена при помощи программы Pathway Tools [Karp et al., Briefings in Bioinformatics, 2010]. Для дополнения исходной аннотации использовалась встроенная в Pathway Tools утилита Hole Filler из пакета PathoLogic.')

document.add_paragraph(u'4) Выравнивание нуклеотидных последовательностей геномов')
document.add_paragraph(u'Построение кольцевой диаграммы было выполнено при помощи программы BRIG [Alikhan et al., BMC Genomics, 2011]. Выравнивание нуклеотидных последовательностей геномов исследуемого и референсных штаммов производилось при помощи утилиты BLASTN.')

document.add_paragraph(u'5) Анализ групп гомологий')
document.add_paragraph(u'Все референсные геномы были скачаны в формате FASTA и проаннотированны утилитой prokka. Поиск ортологов был выполнен с помощью утилиты OrthoMCL. Таблицы получены при помощи разработанных нами скриптов на perl и R.')

document.add_paragraph(u'6) Филогенетический анализ')

document.add_paragraph(u'Для построения филогенетиечского дерева использовались однокопийные гены, представленные у всех рассматриваемых штаммов. Для каждой группы генов было выполнено множественное выравнивание программой ClustalW, затем на основе выравнивания было построено дерево методом Maximal Likelihood при помощи утилит из пакета EMBOSS.')

document.add_paragraph(u'7) Подсчет расстояний между геномными последовательностями')
document.add_paragraph(u'Подсчет расстояний между геномными последовательностями осуществлялся двумя способами: а) по однонуклеотидным заменам в последовательностях общих для всех рассматриваемых геномов генов (core-genome distance); б) по генному составу (признаку наличия-отсутствия генов из пан генома, pan-genome profile).')
document.add_paragraph(u'В первом случае мы применяли утилиту dnadist из пакета EMBOSS для подсчета расстояний по множественному выравниванию каждого из общих генов. В качестве расстояний между геномами использовалось среднее значение расстояний по всем общим генам.')
document.add_paragraph(u'Для подсчета расстояний по генному составу мы вначале составили вектора наличия-отсутствия каждого гена для рассматриваемых геномов. Затем, расстояния на основе метрики L1 (метрика Манхэттена) были получены при помощи скриптов на R.')

document.add_heading(u'Литература', level=2)
document.add_paragraph(u'1. Pevzner, Pavel A., Haixu Tang, and Michael S. Waterman. "An Eulerian path approach to DNA fragment assembly." Proceedings of the National Academy of Sciences 98.17 (2001): 9748-9753.')
document.add_paragraph(u'2. Boetzer, Marten, et al. "Scaffolding pre-assembled contigs using SSPACE." Bioinformatics 27.4 (2011): 578-579.')
document.add_paragraph(u'3. Seemann, Torsten. "Prokka: rapid prokaryotic genome annotation." Bioinformatics (2014): btu153.')
document.add_paragraph(u'4. Karp, Peter D., et al. "Pathway Tools version 13.0: integrated software for pathway/genome informatics and systems biology." Briefings in bioinformatics 11.1 (2010): 40-79.')
document.add_paragraph(u'5. Alikhan, Nabil-Fareed, et al. "BLAST Ring Image Generator (BRIG): simple prokaryote genome comparisons." BMC genomics 12.1 (2011): 402.')

document.save('report.docx')
